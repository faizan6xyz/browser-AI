"""
File-Aware Web Scraping Agent
Handles text, images, PDFs, DOCX, Excel files, and more.
Integrates with browser automation (Playwright) and LLM orchestrators.
"""

import os
import re
import json
import time
import asyncio
from enum import Enum
from typing import Optional, Dict, List, Any, Union
from dataclasses import dataclass, field
from pathlib import Path
from io import BytesIO

import requests
from bs4 import BeautifulSoup


# ============================================================================
# FILE TYPE ENUMERATION
# ============================================================================

class FileType(Enum):
    """Supported file types for scraping"""
    TEXT = "text"
    HTML = "html"
    IMAGE = "image"
    PDF = "pdf"
    DOCX = "docx"
    EXCEL = "excel"
    CSV = "csv"
    UNKNOWN = "unknown"


# MIME type to FileType mapping
MIME_TO_FILETYPE = {
    'text/plain': FileType.TEXT,
    'text/html': FileType.HTML,
    'application/pdf': FileType.PDF,
    'application/msword': FileType.DOCX,
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': FileType.DOCX,
    'application/vnd.ms-excel': FileType.EXCEL,
    'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': FileType.EXCEL,
    'text/csv': FileType.CSV,
    'application/csv': FileType.CSV,
    'image/jpeg': FileType.IMAGE,
    'image/png': FileType.IMAGE,
    'image/gif': FileType.IMAGE,
    'image/webp': FileType.IMAGE,
    'image/svg+xml': FileType.IMAGE,
}

# File extension to FileType mapping
EXT_TO_FILETYPE = {
    '.txt': FileType.TEXT,
    '.html': FileType.HTML,
    '.htm': FileType.HTML,
    '.pdf': FileType.PDF,
    '.doc': FileType.DOCX,
    '.docx': FileType.DOCX,
    '.xls': FileType.EXCEL,
    '.xlsx': FileType.EXCEL,
    '.csv': FileType.CSV,
    '.jpg': FileType.IMAGE,
    '.jpeg': FileType.IMAGE,
    '.png': FileType.IMAGE,
    '.gif': FileType.IMAGE,
    '.webp': FileType.IMAGE,
    '.svg': FileType.IMAGE,
}


# ============================================================================
# DATA CLASSES
# ============================================================================

@dataclass
class FileMetadata:
    """Metadata about a discovered file"""
    url: str
    file_type: FileType
    filename: str
    size_bytes: Optional[int] = None
    mime_type: Optional[str] = None
    source_element: Optional[str] = None  # CSS selector or XPath
    page_url: Optional[str] = None  # URL of the page where file was found


@dataclass
class ProcessedFile:
    """Result of processing a downloaded file"""
    metadata: FileMetadata
    success: bool
    extracted_text: Optional[str] = None
    extracted_data: Optional[Any] = None  # Dict, DataFrame, etc.
    file_path: Optional[str] = None
    error_message: Optional[str] = None
    
    def to_dict(self) -> Dict:
        """Convert to dictionary for LLM/orchestrator consumption"""
        return {
            'url': self.metadata.url,
            'file_type': self.metadata.file_type.value,
            'filename': self.metadata.filename,
            'success': self.success,
            'extracted_text': self.extracted_text[:500] if self.extracted_text else None,  # Truncate for context
            'extracted_data_summary': str(self.extracted_data)[:200] if self.extracted_data else None,
            'error': self.error_message,
        }


# ============================================================================
# FILE TYPE DETECTOR
# ============================================================================

class FileTypeDetector:
    """Detects file types from URLs, headers, and content"""
    
    @staticmethod
    def detect_from_url(url: str) -> FileType:
        """Detect file type from URL extension"""
        try:
            # Remove query parameters
            clean_url = url.split('?')[0]
            ext = '.' + clean_url.split('.')[-1].lower() if '.' in clean_url else ''
            return EXT_TO_FILETYPE.get(ext, FileType.UNKNOWN)
        except:
            return FileType.UNKNOWN
    
    @staticmethod
    def detect_from_mime(mime_type: str) -> FileType:
        """Detect file type from MIME type"""
        if not mime_type:
            return FileType.UNKNOWN
        # Handle charset in MIME type (e.g., "text/html; charset=utf-8")
        mime_base = mime_type.split(';')[0].strip().lower()
        return MIME_TO_FILETYPE.get(mime_base, FileType.UNKNOWN)
    
    @staticmethod
    def detect_from_headers(headers: Dict) -> FileType:
        """Detect file type from HTTP response headers"""
        content_type = headers.get('Content-Type', '')
        return FileTypeDetector.detect_from_mime(content_type)
    
    @staticmethod
    def detect_comprehensive(url: str, headers: Optional[Dict] = None) -> FileType:
        """Comprehensive detection using multiple strategies"""
        # Strategy 1: Check headers first (most reliable)
        if headers:
            file_type = FileTypeDetector.detect_from_headers(headers)
            if file_type != FileType.UNKNOWN:
                return file_type
        
        # Strategy 2: Check URL extension
        file_type = FileTypeDetector.detect_from_url(url)
        if file_type != FileType.UNKNOWN:
            return file_type
        
        return FileType.UNKNOWN


# ============================================================================
# FILE DOWNLOADER
# ============================================================================

class FileDownloader:
    """Downloads files with streaming support and browser automation fallback"""
    
    def __init__(self, download_dir: str = "./downloads", timeout: int = 30):
        self.download_dir = Path(download_dir)
        self.download_dir.mkdir(parents=True, exist_ok=True)
        self.timeout = timeout
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        })
    
    def download_file(self, url: str, filename: Optional[str] = None) -> str:
        """
        Download a file using requests with streaming.
        Returns the local file path.
        """
        if not filename:
            filename = url.split('/')[-1].split('?')[0]
            if not filename or '.' not in filename:
                filename = f"download_{int(time.time())}"
        
        filepath = self.download_dir / filename
        
        try:
            response = self.session.get(url, stream=True, timeout=self.timeout)
            response.raise_for_status()
            
            # Write in chunks to handle large files
            with open(filepath, 'wb') as f:
                for chunk in response.iter_content(chunk_size=8192):
                    if chunk:
                        f.write(chunk)
            
            return str(filepath)
        
        except Exception as e:
            raise Exception(f"Failed to download {url}: {str(e)}")
    
    async def download_with_browser(self, url: str, click_selector: Optional[str] = None) -> str:
        """
        Download file using Playwright browser automation.
        Useful for files behind JavaScript or requiring interaction.
        
        Args:
            url: The page URL containing the download link/button
            click_selector: CSS selector for the download button/link (optional)
        
        Returns:
            Local file path
        """
        try:
            from playwright.async_api import async_playwright
        except ImportError:
            raise ImportError("Install playwright: pip install playwright && playwright install")
        
        async with async_playwright() as p:
            browser = await p.chromium.launch(headless=True)
            page = await browser.new_page()
            
            # Set up download handling
            downloaded_files = []
            
            async def handle_download(download):
                suggested_filename = download.suggested_filename
                filepath = self.download_dir / suggested_filename
                await download.save_as(str(filepath))
                downloaded_files.append(str(filepath))
            
            page.on("download", handle_download)
            
            await page.goto(url, wait_until="networkidle")
            
            # If click selector provided, click it to trigger download
            if click_selector:
                await page.click(click_selector)
                # Wait for download to complete
                await asyncio.sleep(2)
            else:
                # Try to find and click download links automatically
                download_links = await page.query_selector_all('a[href$=".pdf"], a[href$=".docx"], a[href$=".xlsx"]')
                if download_links:
                    await download_links[0].click()
                    await asyncio.sleep(2)
            
            await browser.close()
            
            if downloaded_files:
                return downloaded_files[0]
            else:
                raise Exception("No file was downloaded")
    
    def get_filename_from_url(self, url: str) -> str:
        """Extract filename from URL"""
        filename = url.split('/')[-1].split('?')[0]
        if not filename or '.' not in filename:
            filename = f"file_{int(time.time())}"
        return filename


# ============================================================================
# FILE PROCESSOR
# ============================================================================

class FileProcessor:
    """Processes different file types and extracts content"""
    
    def __init__(self):
        self.processors = {
            FileType.TEXT: self._process_text,
            FileType.HTML: self._process_html,
            FileType.IMAGE: self._process_image,
            FileType.PDF: self._process_pdf,
            FileType.DOCX: self._process_docx,
            FileType.EXCEL: self._process_excel,
            FileType.CSV: self._process_csv,
        }
    
    def process(self, filepath: str, file_type: FileType) -> ProcessedFile:
        """Process a file based on its type"""
        processor = self.processors.get(file_type)
        if not processor:
            return ProcessedFile(
                metadata=FileMetadata(url="", file_type=file_type, filename=filepath),
                success=False,
                error_message=f"No processor for file type: {file_type.value}"
            )
        
        try:
            result = processor(filepath)
            result.success = True
            return result
        except Exception as e:
            return ProcessedFile(
                metadata=FileMetadata(url="", file_type=file_type, filename=filepath),
                success=False,
                error_message=str(e)
            )
    
    def _process_text(self, filepath: str) -> ProcessedFile:
        """Process plain text files"""
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        
        return ProcessedFile(
            metadata=FileMetadata(url="", file_type=FileType.TEXT, filename=Path(filepath).name),
            success=False,  # Will be set by caller
            extracted_text=text
        )
    
    def _process_html(self, filepath: str) -> ProcessedFile:
        """Process HTML files - extract text content"""
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            html_content = f.read()
        
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        text = soup.get_text(separator='\n')
        # Clean up whitespace
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        clean_text = '\n'.join(lines)
        
        return ProcessedFile(
            metadata=FileMetadata(url="", file_type=FileType.HTML, filename=Path(filepath).name),
            success=False,
            extracted_text=clean_text
        )
    
    def _process_image(self, filepath: str) -> ProcessedFile:
        """Process image files - optional OCR"""
        try:
            from PIL import Image
            img = Image.open(filepath)
            
            metadata_info = {
                'format': img.format,
                'size': img.size,
                'mode': img.mode,
            }
            
            # Try OCR if pytesseract is available
            extracted_text = None
            try:
                import pytesseract
                extracted_text = pytesseract.image_to_string(img)
            except ImportError:
                pass  # OCR not available
            
            return ProcessedFile(
                metadata=FileMetadata(url="", file_type=FileType.IMAGE, filename=Path(filepath).name),
                success=False,
                extracted_text=extracted_text,
                extracted_data=metadata_info
            )
        
        except Exception as e:
            return ProcessedFile(
                metadata=FileMetadata(url="", file_type=FileType.IMAGE, filename=Path(filepath).name),
                success=False,
                error_message=f"Failed to process image: {str(e)}"
            )
    
    def _process_pdf(self, filepath: str) -> ProcessedFile:
        """Process PDF files - extract text"""
        try:
            import PyPDF2
            
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                
                text_parts = []
                for page_num in range(len(reader.pages)):
                    page = reader.pages[page_num]
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                
                full_text = '\n\n'.join(text_parts)
                
                metadata_info = {
                    'num_pages': len(reader.pages),
                }
                
                return ProcessedFile(
                    metadata=FileMetadata(url="", file_type=FileType.PDF, filename=Path(filepath).name),
                    success=False,
                    extracted_text=full_text,
                    extracted_data=metadata_info
                )
        
        except ImportError:
            return ProcessedFile(
                metadata=FileMetadata(url="", file_type=FileType.PDF, filename=Path(filepath).name),
                success=False,
                error_message="PyPDF2 not installed. Install with: pip install PyPDF2"
            )
        except Exception as e:
            return ProcessedFile(
                metadata=FileMetadata(url="", file_type=FileType.PDF, filename=Path(filepath).name),
                success=False,
                error_message=f"Failed to process PDF: {str(e)}"
            )
    
    def _process_docx(self, filepath: str) -> ProcessedFile:
        """Process DOCX files - extract text and tables"""
        try:
            from docx import Document
            
            doc = Document(filepath)
            
            # Extract paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Extract tables
            tables_data = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables_data.append(table_data)
            
            full_text = '\n\n'.join(paragraphs)
            
            return ProcessedFile(
                metadata=FileMetadata(url="", file_type=FileType.DOCX, filename=Path(filepath).name),
                success=False,
                extracted_text=full_text,
                extracted_data={'tables': tables_data}
            )
        
        except ImportError:
            return ProcessedFile(
                metadata=FileMetadata(url="", file_type=FileType.DOCX, filename=Path(filepath).name),
                success=False,
                error_message="python-docx not installed. Install with: pip install python-docx"
            )
        except Exception as e:
            return ProcessedFile(
                metadata=FileMetadata(url="", file_type=FileType.DOCX, filename=Path(filepath).name),
                success=False,
                error_message=f"Failed to process DOCX: {str(e)}"
            )
    
    def _process_excel(self, filepath: str) -> ProcessedFile:
        """Process Excel files - extract structured data"""
        try:
            import pandas as pd
            
            # Try different engines
            try:
                df = pd.read_excel(filepath, engine='openpyxl')
            except:
                df = pd.read_excel(filepath, engine='xlrd')
            
            # Convert to dict for serialization
            data_dict = df.to_dict('records')
            
            # Also create text representation
            text_repr = df.to_string()
            
            metadata_info = {
                'num_rows': len(df),
                'num_columns': len(df.columns),
                'columns': list(df.columns),
            }
            
            return ProcessedFile(
                metadata=FileMetadata(url="", file_type=FileType.EXCEL, filename=Path(filepath).name),
                success=False,
                extracted_text=text_repr,
                extracted_data={'data': data_dict, 'metadata': metadata_info}
            )
        
        except ImportError:
            return ProcessedFile(
                metadata=FileMetadata(url="", file_type=FileType.EXCEL, filename=Path(filepath).name),
                success=False,
                error_message="pandas/openpyxl not installed. Install with: pip install pandas openpyxl"
            )
        except Exception as e:
            return ProcessedFile(
                metadata=FileMetadata(url="", file_type=FileType.EXCEL, filename=Path(filepath).name),
                success=False,
                error_message=f"Failed to process Excel: {str(e)}"
            )
    
    def _process_csv(self, filepath: str) -> ProcessedFile:
        """Process CSV files"""
        try:
            import pandas as pd
            
            df = pd.read_csv(filepath)
            
            data_dict = df.to_dict('records')
            text_repr = df.to_string()
            
            metadata_info = {
                'num_rows': len(df),
                'num_columns': len(df.columns),
                'columns': list(df.columns),
            }
            
            return ProcessedFile(
                metadata=FileMetadata(url="", file_type=FileType.CSV, filename=Path(filepath).name),
                success=False,
                extracted_text=text_repr,
                extracted_data={'data': data_dict, 'metadata': metadata_info}
            )
        
        except ImportError:
            return ProcessedFile(
                metadata=FileMetadata(url="", file_type=FileType.CSV, filename=Path(filepath).name),
                success=False,
                error_message="pandas not installed. Install with: pip install pandas"
            )
        except Exception as e:
            return ProcessedFile(
                metadata=FileMetadata(url="", file_type=FileType.CSV, filename=Path(filepath).name),
                success=False,
                error_message=f"Failed to process CSV: {str(e)}"
            )


# ============================================================================
# WEB SCRAPER AGENT
# ============================================================================

class WebScraperAgent:
    """
    Main agent for web scraping with multi-format file support.
    Discovers, downloads, and processes various file types from web pages.
    """
    
    def __init__(self, download_dir: str = "./downloads"):
        self.detector = FileTypeDetector()
        self.downloader = FileDownloader(download_dir=download_dir)
        self.processor = FileProcessor()
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        })
    
    def discover_files_on_page(self, url: str) -> List[FileMetadata]:
        """
        Discover all downloadable files on a webpage.
        
        Args:
            url: The webpage URL to scan
        
        Returns:
            List of FileMetadata objects for discovered files
        """
        try:
            response = self.session.get(url, timeout=10)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.text, 'html.parser')
            files = []
            
            # Strategy 1: Find links with file extensions
            for link in soup.find_all('a', href=True):
                href = link['href']
                
                # Skip relative URLs that don't look like files
                if href.startswith('#') or href.startswith('javascript:'):
                    continue
                
                # Make absolute URL
                if href.startswith('/'):
                    from urllib.parse import urljoin
                    href = urljoin(url, href)
                
                # Detect file type
                file_type = self.detector.detect_from_url(href)
                
                if file_type != FileType.UNKNOWN:
                    filename = self.downloader.get_filename_from_url(href)
                    files.append(FileMetadata(
                        url=href,
                        file_type=file_type,
                        filename=filename,
                        source_element=f"a[href='{href}']",
                        page_url=url
                    ))
            
            # Strategy 2: Find img tags
            for img in soup.find_all('img', src=True):
                src = img['src']
                
                if src.startswith('data:'):  # Skip base64 images
                    continue
                
                if src.startswith('/'):
                    from urllib.parse import urljoin
                    src = urljoin(url, src)
                
                file_type = self.detector.detect_from_url(src)
                if file_type == FileType.UNKNOWN:
                    file_type = FileType.IMAGE  # Assume image if in img tag
                
                filename = self.downloader.get_filename_from_url(src)
                files.append(FileMetadata(
                    url=src,
                    file_type=file_type,
                    filename=filename,
                    source_element=f"img[src='{src}']",
                    page_url=url
                ))
            
            # Strategy 3: Find embed/object tags (often PDFs)
            for embed in soup.find_all(['embed', 'object']):
                src = embed.get('src') or embed.get('data')
                if src:
                    if src.startswith('/'):
                        from urllib.parse import urljoin
                        src = urljoin(url, src)
                    
                    file_type = self.detector.detect_from_url(src)
                    if file_type == FileType.UNKNOWN:
                        file_type = FileType.PDF  # Common for embeds
                    
                    filename = self.downloader.get_filename_from_url(src)
                    files.append(FileMetadata(
                        url=src,
                        file_type=file_type,
                        filename=filename,
                        source_element=f"embed[src='{src}']",
                        page_url=url
                    ))
            
            return files
        
        except Exception as e:
            print(f"Error discovering files on {url}: {str(e)}")
            return []
    
    def download_and_process(self, file_metadata: FileMetadata) -> ProcessedFile:
        """
        Download and process a single file.
        
        Args:
            file_metadata: Metadata about the file to download
        
        Returns:
            ProcessedFile with extracted content
        """
        try:
            # Download the file
            filepath = self.downloader.download_file(
                file_metadata.url,
                filename=file_metadata.filename
            )
            
            # Update metadata with actual file info
            file_metadata.file_path = filepath
            if os.path.exists(filepath):
                file_metadata.size_bytes = os.path.getsize(filepath)
            
            # Process the file
            result = self.processor.process(filepath, file_metadata.file_type)
            result.metadata = file_metadata
            result.file_path = filepath
            
            return result
        
        except Exception as e:
            return ProcessedFile(
                metadata=file_metadata,
                success=False,
                error_message=f"Failed to download/process: {str(e)}"
            )
    
    async def scrape_page(self, url: str, file_types: Optional[List[FileType]] = None) -> List[ProcessedFile]:
        """
        Main method: Scrape a webpage, discover files, download and process them.
        
        Args:
            url: The webpage URL
            file_types: Optional filter for specific file types
        
        Returns:
            List of ProcessedFile results
        """
        print(f"🔍 Discovering files on: {url}")
        
        # Discover files
        files = self.discover_files_on_page(url)
        
        # Filter by file types if specified
        if file_types:
            files = [f for f in files if f.file_type in file_types]
        
        print(f"📁 Found {len(files)} files")
        
        # Download and process each file
        results = []
        for i, file_meta in enumerate(files, 1):
            print(f"  [{i}/{len(files)}] Processing: {file_meta.filename} ({file_meta.file_type.value})")
            result = self.download_and_process(file_meta)
            results.append(result)
            
            if result.success:
                print(f"    ✅ Success")
            else:
                print(f"    ❌ Failed: {result.error_message}")
        
        return results
    
    def scrape_single_file(self, url: str) -> ProcessedFile:
        """
        Download and process a single file directly (not from a webpage).
        
        Args:
            url: Direct URL to the file
        
        Returns:
            ProcessedFile result
        """
        print(f"📥 Downloading: {url}")
        
        # Detect file type
        try:
            response = self.session.head(url, timeout=5, allow_redirects=True)
            file_type = self.detector.detect_comprehensive(url, response.headers)
        except:
            file_type = self.detector.detect_from_url(url)
        
        if file_type == FileType.UNKNOWN:
            file_type = FileType.TEXT  # Default fallback
        
        filename = self.downloader.get_filename_from_url(url)
        
        metadata = FileMetadata(
            url=url,
            file_type=file_type,
            filename=filename
        )
        
        return self.download_and_process(metadata)


# ============================================================================
# LLM ORCHESTRATOR INTEGRATION
# ============================================================================

class LLMAgentOrchestrator:
    """
    Example integration with LLM-based agent orchestrator.
    Shows how to structure prompts and parse actions for file scraping.
    """
    
    def __init__(self, scraper_agent: WebScraperAgent):
        self.scraper = scraper_agent
    
    def generate_scraping_prompt(self, user_intent: str, page_context: Dict) -> str:
        """
        Generate a prompt for an LLM to decide what files to scrape.
        
        Args:
            user_intent: What the user wants to extract
            page_context: Information about the current page
        
        Returns:
            Formatted prompt string
        """
        prompt = f"""
You are a web scraping agent with multi-format file support.

USER INTENT: "{user_intent}"

PAGE CONTEXT:
- URL: {page_context.get('url', 'N/A')}
- Title: {page_context.get('title', 'N/A')}
- Available files:
{json.dumps(page_context.get('available_files', []), indent=2)}

SUPPORTED FILE TYPES:
- Images (JPG, PNG, GIF, WebP) - can extract text via OCR
- PDFs - extract text content
- DOCX - extract text and tables
- Excel (XLSX/XLS) - extract structured data
- CSV - extract tabular data
- Text/HTML - extract readable text

AVAILABLE ACTIONS:
1. DISCOVER_FILES(url) - Scan a webpage for downloadable files
2. DOWNLOAD_FILE(url, filename) - Download a specific file
3. PROCESS_FILE(filepath, file_type) - Extract content from a file
4. SCRAPE_PAGE(url, file_types) - Full pipeline: discover, download, process

INSTRUCTIONS:
Based on the user intent and available files, determine:
1. Which files are relevant to the user's goal?
2. What actions should be taken?
3. Return a JSON array of actions to execute.

Example output format:
{{
  "actions": [
    {{"action": "SCRAPE_PAGE", "params": {{"url": "...", "file_types": ["pdf", "excel"]}}}},
    {{"action": "DOWNLOAD_FILE", "params": {{"url": "...", "filename": "report.pdf"}}}}
  ],
  "reasoning": "Brief explanation of why these actions were chosen"
}}

Return ONLY valid JSON.
"""
        return prompt
    
    def parse_llm_response(self, llm_output: str) -> List[Dict]:
        """
        Parse LLM response to extract actions.
        
        Args:
            llm_output: Raw LLM output (should be JSON)
        
        Returns:
            List of action dictionaries
        """
        try:
            # Try to extract JSON from the response
            if '{' in llm_output:
                json_start = llm_output.index('{')
                json_end = llm_output.rindex('}') + 1
                json_str = llm_output[json_start:json_end]
                data = json.loads(json_str)
                return data.get('actions', [])
            else:
                return []
        except Exception as e:
            print(f"Failed to parse LLM response: {e}")
            return []
    
    async def execute_actions(self, actions: List[Dict]) -> List[ProcessedFile]:
        """
        Execute actions returned by LLM.
        
        Args:
            actions: List of action dictionaries
        
        Returns:
            List of ProcessedFile results
        """
        results = []
        
        for action in actions:
            action_type = action.get('action')
            params = action.get('params', {})
            
            print(f"⚙️  Executing: {action_type}")
            
            if action_type == "SCRAPE_PAGE":
                url = params.get('url')
                file_types_str = params.get('file_types', [])
                file_types = [FileType(ft) for ft in file_types_str if ft in [e.value for e in FileType]]
                
                if url:
                    page_results = await self.scraper.scrape_page(url, file_types if file_types else None)
                    results.extend(page_results)
            
            elif action_type == "DOWNLOAD_FILE":
                url = params.get('url')
                if url:
                    result = self.scraper.scrape_single_file(url)
                    results.append(result)
            
            elif action_type == "DISCOVER_FILES":
                url = params.get('url')
                if url:
                    files = self.scraper.discover_files_on_page(url)
                    print(f"Discovered {len(files)} files")
                    # Could add logic to auto-download discovered files
            
            else:
                print(f"Unknown action: {action_type}")
        
        return results


# ============================================================================
# EXAMPLE USAGE
# ============================================================================

async def example_basic_usage():
    """Basic example: Scrape a page and process all files"""
    print("=" * 80)
    print("EXAMPLE 1: Basic File Scraping")
    print("=" * 80)
    
    agent = WebScraperAgent(download_dir="./downloads")
    
    # Example: Scrape a page (replace with your target URL)
    # Note: This is a demo - use a real URL with actual files
    url = "https://www.example.com"  # Replace with actual URL
    
    try:
        results = await agent.scrape_page(url)
        
        print(f"\n📊 Results Summary:")
        print(f"   Total files processed: {len(results)}")
        print(f"   Successful: {sum(1 for r in results if r.success)}")
        print(f"   Failed: {sum(1 for r in results if not r.success)}")
        
        for result in results:
            if result.success:
                print(f"\n✅ {result.metadata.filename}")
                print(f"   Type: {result.metadata.file_type.value}")
                if result.extracted_text:
                    preview = result.extracted_text[:200].replace('\n', ' ')
                    print(f"   Preview: {preview}...")
    
    except Exception as e:
        print(f"Error: {e}")


async def example_llm_orchestration():
    """Example: Use with LLM orchestrator"""
    print("\n" + "=" * 80)
    print("EXAMPLE 2: LLM Orchestrator Integration")
    print("=" * 80)
    
    agent = WebScraperAgent(download_dir="./downloads")
    orchestrator = LLMAgentOrchestrator(agent)
    
    # Simulate user intent
    user_intent = "Extract all financial reports and data tables from this page"
    
    # Simulate page context (in real scenario, this comes from browser automation)
    page_context = {
        'url': 'https://example.com/reports',
        'title': 'Financial Reports',
        'available_files': [
            {'url': 'https://example.com/report.pdf', 'type': 'pdf'},
            {'url': 'https://example.com/data.xlsx', 'type': 'excel'},
        ]
    }
    
    # Generate prompt for LLM
    prompt = orchestrator.generate_scraping_prompt(user_intent, page_context)
    print("📝 Generated Prompt (first 500 chars):")
    print(prompt[:500] + "...")
    
    # Simulate LLM response (in real scenario, call your LLM API)
    simulated_llm_response = '''
    {
      "actions": [
        {
          "action": "DOWNLOAD_FILE",
          "params": {
            "url": "https://example.com/report.pdf",
            "filename": "report.pdf"
          }
        },
        {
          "action": "DOWNLOAD_FILE",
          "params": {
            "url": "https://example.com/data.xlsx",
            "filename": "data.xlsx"
          }
        }
      ],
      "reasoning": "User wants financial reports and data tables. PDF likely contains report text, Excel contains structured data."
    }
    '''
    
    # Parse and execute actions
    actions = orchestrator.parse_llm_response(simulated_llm_response)
    print(f"\n🎯 Parsed {len(actions)} actions from LLM")
    
    # Execute (commented out since URLs are fake)
    # results = await orchestrator.execute_actions(actions)


def example_single_file():
    """Example: Process a single file directly"""
    print("\n" + "=" * 80)
    print("EXAMPLE 3: Single File Processing")
    print("=" * 80)
    
    agent = WebScraperAgent(download_dir="./downloads")
    
    # Example: Process a single file URL
    # Replace with actual file URL
    file_url = "https://example.com/document.pdf"
    
    try:
        result = agent.scrape_single_file(file_url)
        
        if result.success:
            print(f"✅ Successfully processed: {result.metadata.filename}")
            print(f"   Type: {result.metadata.file_type.value}")
            if result.extracted_text:
                print(f"   Text length: {len(result.extracted_text)} characters")
        else:
            print(f"❌ Failed: {result.error_message}")
    
    except Exception as e:
        print(f"Error: {e}")


# ============================================================================
# MAIN ENTRY POINT
# ============================================================================

if __name__ == "__main__":
    print("🤖 File-Aware Web Scraping Agent")
    print("=" * 80)
    print("\nThis agent can scrape and process:")
    print("  • Images (with optional OCR)")
    print("  • PDFs (text extraction)")
    print("  • DOCX files (text and tables)")
    print("  • Excel files (structured data)")
    print("  • CSV files (tabular data)")
    print("  • Text/HTML files")
    print("\n" + "=" * 80)
    
    # Run examples
    asyncio.run(example_basic_usage())
    asyncio.run(example_llm_orchestration())
    example_single_file()
    
    print("\n" + "=" * 80)
    print("💡 Tips:")
    print("  1. Replace example URLs with your target URLs")
    print("  2. Install optional dependencies for full features:")
    print("     pip install pytesseract (for OCR)")
    print("     pip install playwright && playwright install (for browser automation)")
    print("  3. Integrate with your LLM orchestrator using LLMAgentOrchestrator class")
    print("=" * 80)